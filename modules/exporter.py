"""exporter.py v4.3 - Full daily Excel with vocab+notes, beautiful formatting"""
import io, urllib.parse
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Styles ────────────────────────────────────────────────────────────────────
TOPIC_LABELS = {'finance':'Tài chính','politics':'Chính trị','tech_ai':'AI & Tech','general':'Chung'}
IMP_LABELS = {1:'Thấp',2:'Bình thường',3:'Đáng đọc',4:'Quan trọng',5:'Rất quan trọng'}

THIN = Side(style='thin', color='D0D7E1')
MED = Side(style='medium', color='8899AA')
BD = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
BD_M = Border(left=MED, right=MED, top=MED, bottom=MED)

# Header fills
H_INTL = PatternFill('solid', fgColor='1E40AF')
H_VN = PatternFill('solid', fgColor='065F46')
H_VOCAB = PatternFill('solid', fgColor='5B21B6')
H_NOTE = PatternFill('solid', fgColor='92400E')
H_STAT = PatternFill('solid', fgColor='1E293B')
H_DARK = PatternFill('solid', fgColor='0F172A')
H_DARK2 = PatternFill('solid', fgColor='1E293B')

# Row fills
R_EVEN = PatternFill('solid', fgColor='F8FAFC')
R_ODD = PatternFill('solid', fgColor='FFFFFF')
R_FIN = PatternFill('solid', fgColor='ECFDF5')
R_POL = PatternFill('solid', fgColor='F5F3FF')
R_TEC = PatternFill('solid', fgColor='FFF7ED')
TOPIC_FILL = {'finance':R_FIN, 'politics':R_POL, 'tech_ai':R_TEC, 'general':R_ODD}
TOPIC_CLR = {'finance':'065F46', 'politics':'4C1D95', 'tech_ai':'9A3412', 'general':'374151'}
IMP_CLR = {5:'991B1B', 4:'DC2626', 3:'D97706', 2:'6B7280', 1:'9CA3AF'}

def _reader_url(url, port=8765, cat='international'):
    if not url or not url.startswith('http'): return ''
    return f"http://localhost:{port}/reader?url={urllib.parse.quote(url, safe='')}&cat={cat}"

def _banner(ws, title, subtitle, row, ncols, fill=H_DARK):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    c = ws.cell(row=row, column=1, value=title)
    c.font = Font(bold=True, size=16, color='FFFFFF', name='Calibri')
    c.fill = fill
    c.alignment = Alignment(horizontal='center', vertical='center')
    c.border = BD_M
    ws.row_dimensions[row].height = 36
    ws.merge_cells(start_row=row+1, start_column=1, end_row=row+1, end_column=ncols)
    c2 = ws.cell(row=row+1, column=1, value=subtitle)
    c2.font = Font(italic=True, size=10, color='94A3B8', name='Calibri')
    c2.fill = H_DARK2
    c2.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[row+1].height = 18
    return row + 2

def _header(ws, row, specs, fill):
    for i, (val, w) in enumerate(specs, 1):
        c = ws.cell(row=row, column=i, value=val)
        c.font = Font(bold=True, color='FFFFFF', size=10, name='Calibri')
        c.fill = fill
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        c.border = BD
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.row_dimensions[row].height = 28

def _c(ws, row, col, val, fill=None, font=None, link=None, halign='left'):
    c = ws.cell(row=row, column=col, value=str(val) if val is not None else '')
    c.font = font or Font(size=10, name='Calibri')
    c.alignment = Alignment(wrap_text=True, vertical='top', horizontal=halign)
    c.border = BD
    if fill: c.fill = fill
    if link and str(link).startswith('http'):
        c.hyperlink = link
        c.font = Font(color='1D4ED8', underline='single', size=10, name='Calibri')
    return c

# ── Main daily export (articles + vocab + notes + stats) ──────────────────────
def export_excel_day(articles, notes, date, port=8765, vocab=None):
    wb = Workbook()
    intl = [a for a in articles if a['category'] == 'international']
    viet = [a for a in articles if a['category'] == 'vietnam']
    vocab = vocab or []

    # ── Sheet 1: International ────────────────────────────────────────────
    ws = wb.active
    ws.title = '🌍 Quốc tế'
    ws.freeze_panes = 'A4'
    ws.sheet_properties.tabColor = '1E40AF'
    _banner(ws, f'📰 TIN QUỐC TẾ — {date}',
            f'{len(intl)} bài · Yahoo Finance, The Guardian, DW, CNBC, TechCrunch, Wired', 1, 9, H_INTL)
    _header(ws, 3, [('#',4),('Nguồn',14),('Chủ đề',12),('Tiêu đề gốc',40),('Mô tả',52),
                     ('Tiêu đề VN',36),('★',8),('Link',12),('Đọc',12)], H_INTL)
    for ri, a in enumerate(intl, 1):
        row = ri + 3
        base = R_EVEN if ri%2==0 else R_ODD
        topic = a.get('topic','general')
        imp = a.get('importance',3)
        stars = '★'*imp + '☆'*(5-imp)
        _c(ws,row,1,ri,base,Font(bold=True,size=10,name='Calibri',color='374151'),halign='center')
        _c(ws,row,2,a.get('source',''),base,Font(size=10,name='Calibri',color='374151'))
        tc = ws.cell(row=row,column=3,value=TOPIC_LABELS.get(topic,''))
        tc.font = Font(bold=True,size=10,name='Calibri',color=TOPIC_CLR.get(topic,'374151'))
        tc.fill = TOPIC_FILL.get(topic,R_ODD); tc.alignment = Alignment(wrap_text=True,vertical='top'); tc.border = BD
        _c(ws,row,4,a.get('title',''),base)
        _c(ws,row,5,a.get('description',''),base,Font(size=9,name='Calibri',color='64748B'))
        _c(ws,row,6,a.get('title_vi',''),base,Font(size=10,name='Calibri',color='1E40AF'))
        _c(ws,row,7,stars,base,Font(size=11,name='Calibri',color=IMP_CLR.get(imp,'6B7280')),halign='center')
        _c(ws,row,8,'Mở bài ↗',base,link=a.get('url',''))
        _c(ws,row,9,'Song ngữ ↗',base,link=_reader_url(a.get('url',''),port,'international'))
        ws.row_dimensions[row].height = 72

    # ── Sheet 2: Vietnam ──────────────────────────────────────────────────
    ws2 = wb.create_sheet('🇻🇳 Việt Nam')
    ws2.freeze_panes = 'A4'
    ws2.sheet_properties.tabColor = '065F46'
    _banner(ws2, f'📰 TIN VIỆT NAM — {date}',
            f'{len(viet)} bài · VnExpress, CafeF, Tuổi Trẻ, Genk', 1, 8, H_VN)
    _header(ws2, 3, [('#',4),('Nguồn',14),('Chủ đề',12),('Tiêu đề',44),('Mô tả',56),
                      ('★',8),('Link',12),('Đọc',12)], H_VN)
    for ri, a in enumerate(viet, 1):
        row = ri + 3
        base = R_EVEN if ri%2==0 else R_ODD
        topic = a.get('topic','general')
        imp = a.get('importance',3)
        stars = '★'*imp + '☆'*(5-imp)
        _c(ws2,row,1,ri,base,Font(bold=True,size=10,name='Calibri'),halign='center')
        _c(ws2,row,2,a.get('source',''),base)
        tc = ws2.cell(row=row,column=3,value=TOPIC_LABELS.get(topic,''))
        tc.font = Font(bold=True,size=10,name='Calibri',color=TOPIC_CLR.get(topic,'374151'))
        tc.fill = TOPIC_FILL.get(topic,R_ODD); tc.alignment = Alignment(wrap_text=True,vertical='top'); tc.border = BD
        _c(ws2,row,4,a.get('title',''),base)
        _c(ws2,row,5,a.get('description',''),base,Font(size=9,name='Calibri',color='64748B'))
        _c(ws2,row,6,stars,base,Font(size=11,name='Calibri',color=IMP_CLR.get(imp,'6B7280')),halign='center')
        _c(ws2,row,7,'Mở bài ↗',base,link=a.get('url',''))
        _c(ws2,row,8,'Đọc bài ↗',base,link=_reader_url(a.get('url',''),port,'vietnam'))
        ws2.row_dimensions[row].height = 72

    # ── Sheet 3: Vocabulary ───────────────────────────────────────────────
    ws3 = wb.create_sheet('📝 Từ vựng')
    ws3.freeze_panes = 'A4'
    ws3.sheet_properties.tabColor = '5B21B6'
    _banner(ws3, f'📝 TỪ VỰNG — {date}',
            f'{len(vocab)} từ/cụm từ đã lưu', 1, 7, H_VOCAB)
    _header(ws3, 3, [('#',4),('Loại',10),('Văn bản gốc',36),('Bản dịch',36),('Bài viết nguồn',40),('Link',12),('Đọc SN',12)], H_VOCAB)
    R_NOTE = PatternFill('solid', fgColor='FEF2F2')
    for ri, v in enumerate(vocab, 1):
        row = ri + 3
        is_note = v.get('type','translate') == 'note'
        base = R_NOTE if is_note else (R_EVEN if ri%2==0 else R_ODD)
        _c(ws3,row,1,ri,base,Font(bold=True,size=10,name='Calibri'),halign='center')
        type_label = '📝 Ghi chú' if is_note else '📖 Dịch'
        type_clr = 'DC2626' if is_note else '065F46'
        _c(ws3,row,2,type_label,base,Font(bold=True,size=10,name='Calibri',color=type_clr))
        _c(ws3,row,3,v.get('original_text',''),base,Font(size=11,name='Calibri',color='DC2626' if is_note else '1E293B'))
        _c(ws3,row,4,v.get('translated_text',''),base,Font(size=11,name='Calibri',color='1E40AF'))
        _c(ws3,row,5,v.get('article_title',''),base,Font(size=9,name='Calibri',color='64748B'))
        _c(ws3,row,6,'Mở bài ↗',base,link=v.get('article_url',''))
        _c(ws3,row,7,'Song ngữ ↗',base,link=_reader_url(v.get('article_url',''),port))
        ws3.row_dimensions[row].height = 32
    if not vocab:
        ws3.merge_cells('A4:G4')
        _c(ws3,4,1,'(Chưa có từ vựng)',R_EVEN,Font(italic=True,size=11,color='94A3B8',name='Calibri'),halign='center')

    # ── Sheet 4: Notes ────────────────────────────────────────────────────
    ws4 = wb.create_sheet('📒 Ghi chú')
    ws4.sheet_properties.tabColor = '92400E'
    _banner(ws4, f'📒 GHI CHÚ — {date}', '', 1, 2, H_NOTE)
    ws4.column_dimensions['A'].width = 20
    ws4.column_dimensions['B'].width = 80
    note_text = notes or ''
    if not note_text.strip():
        ws4.merge_cells('A3:B3')
        c = ws4.cell(row=3, column=1, value='(Chưa có ghi chú)')
        c.font = Font(italic=True, size=11, color='94A3B8', name='Calibri')
        c.alignment = Alignment(wrap_text=True, vertical='top')
    else:
        # Parse notes: split by 📌 markers for structured display
        import re as _re
        note_entries = _re.split(r'(📌\s*\[.*?\])', note_text)
        nr = 3
        R_NOTE_BG = PatternFill('solid', fgColor='FEF2F2')
        for part in note_entries:
            part = part.strip()
            if not part: continue
            if part.startswith('📌'):
                # Article title marker
                _c(ws4, nr, 1, part, R_NOTE_BG, Font(bold=True, size=11, name='Calibri', color='92400E'))
                _c(ws4, nr, 2, '', R_NOTE_BG)
                ws4.row_dimensions[nr].height = 24
                nr += 1
            else:
                # Note content - check for quoted text
                lines = part.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line: continue
                    is_quote = line.startswith('"') and line.endswith('"')
                    if is_quote:
                        _c(ws4, nr, 1, '🔴', R_NOTE_BG, halign='center')
                        _c(ws4, nr, 2, line, R_NOTE_BG, Font(size=11, name='Calibri', color='DC2626', bold=True))
                    else:
                        _c(ws4, nr, 1, '', R_ODD)
                        _c(ws4, nr, 2, line, R_ODD, Font(size=11, name='Calibri', color='1E293B'))
                    ws4.row_dimensions[nr].height = 28
                    nr += 1

    # ── Sheet 5: Dashboard ────────────────────────────────────────────────
    ws5 = wb.create_sheet('📊 Thống kê')
    ws5.sheet_properties.tabColor = '1E293B'
    _banner(ws5, '📊 THỐNG KÊ TỔNG HỢP', date, 1, 3, H_STAT)
    stats = [
        ('📅 Ngày', date, ''),
        ('📰 Tổng bài viết', len(articles), ''),
        ('🌍 Quốc tế', len(intl), f'{len(intl)/max(len(articles),1)*100:.0f}%'),
        ('🇻🇳 Việt Nam', len(viet), f'{len(viet)/max(len(articles),1)*100:.0f}%'),
        ('', '', ''),
        ('💰 Tài chính', sum(1 for a in articles if a.get('topic')=='finance'), ''),
        ('🏛️ Chính trị', sum(1 for a in articles if a.get('topic')=='politics'), ''),
        ('🤖 AI & Tech', sum(1 for a in articles if a.get('topic')=='tech_ai'), ''),
        ('📋 Chung', sum(1 for a in articles if a.get('topic','general')=='general'), ''),
        ('', '', ''),
        ('🔥 Rất quan trọng (★★★★★)', sum(1 for a in articles if a.get('importance',3)>=5), ''),
        ('⭐ Quan trọng (★★★★☆)', sum(1 for a in articles if a.get('importance',3)==4), ''),
        ('📖 Đáng đọc (★★★☆☆)', sum(1 for a in articles if a.get('importance',3)==3), ''),
        ('', '', ''),
        ('📝 Từ vựng đã lưu', len(vocab), ''),
        ('📒 Có ghi chú', 'Có' if notes else 'Không', ''),
        ('', '', ''),
        ('🕐 Xuất lúc', datetime.now().strftime('%d/%m/%Y %H:%M:%S'), ''),
    ]
    for ri, (k, v, pct) in enumerate(stats, 3):
        base = R_EVEN if ri%2==0 else R_ODD
        if not k:  # Empty separator row
            ws5.row_dimensions[ri].height = 8
            continue
        _c(ws5,ri,1,k,base,Font(bold=True,size=11,name='Calibri'))
        _c(ws5,ri,2,str(v),base,Font(size=11,name='Calibri',color='1E293B'))
        if pct:
            _c(ws5,ri,3,pct,base,Font(size=10,name='Calibri',color='64748B'))
        ws5.row_dimensions[ri].height = 24
    ws5.column_dimensions['A'].width = 32
    ws5.column_dimensions['B'].width = 20
    ws5.column_dimensions['C'].width = 12

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

# ── Word export ───────────────────────────────────────────────────────────────
def export_word_day(articles, notes, date, vocab=None):
    vocab = vocab or []
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)

    h = doc.add_heading(f'BẢN TIN NGÀY {date}', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    doc.add_paragraph()

    for cat, lbl, rgb in [('international','🌍 TIN QUỐC TẾ',RGBColor(0x1E,0x40,0xAF)),
                           ('vietnam','🇻🇳 TIN VIỆT NAM',RGBColor(0x06,0x5F,0x46))]:
        arts = [a for a in articles if a['category'] == cat]
        if not arts: continue
        h2 = doc.add_heading(lbl, 1)
        h2.runs[0].font.color.rgb = rgb
        by_topic = {}
        for a in arts: by_topic.setdefault(a.get('topic','general'), []).append(a)
        for topic, items in by_topic.items():
            doc.add_heading(f'  {TOPIC_LABELS.get(topic,topic)}', 2)
            for a in sorted(items, key=lambda x: -x.get('importance',3)):
                imp = a.get('importance',3)
                stars = '★'*imp + '☆'*(5-imp)
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(f'[{stars}] {a.get("title_vi") or a.get("title","")}')
                r.bold = True; r.font.size = Pt(11)
                desc = a.get('description_vi') or a.get('description','')
                if desc:
                    dp = doc.add_paragraph(f'    {desc}')
                    dp.runs[0].font.size = Pt(10)
                    dp.runs[0].font.color.rgb = RGBColor(0x64,0x74,0x8B)
                if a.get('url'):
                    u = doc.add_paragraph(f'    {a["url"]}')
                    u.runs[0].font.size = Pt(8)
                    u.runs[0].font.color.rgb = RGBColor(0x1E,0x40,0xAF)

    # Vocabulary section
    if vocab:
        doc.add_heading('📝 TỪ VỰNG', 1)
        for v in vocab:
            p = doc.add_paragraph(style='List Bullet')
            r1 = p.add_run(v.get('original_text',''))
            r1.bold = True; r1.font.size = Pt(11)
            r2 = p.add_run(f'  →  {v.get("translated_text","")}')
            r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x1E,0x40,0xAF)

    # Notes section
    if notes:
        doc.add_heading('📒 GHI CHÚ', 1)
        doc.add_paragraph(notes)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# ── Vocab-only export ─────────────────────────────────────────────────────────
def export_vocabulary_excel(vocab, port=8765):
    wb = Workbook()
    ws = wb.active
    ws.title = 'Từ vựng'
    ws.freeze_panes = 'A4'
    _banner(ws, 'TỪ VỰNG ĐÃ LƯU', f'{len(vocab)} mục · {datetime.now().strftime("%d/%m/%Y %H:%M")}', 1, 5, H_VOCAB)
    _header(ws, 3, [('#',4),('Văn bản gốc',40),('Bản dịch',40),('Bài viết nguồn',44),('Link',12)], H_VOCAB)
    for ri, v in enumerate(vocab, 1):
        row = ri + 3
        base = R_EVEN if ri%2==0 else R_ODD
        _c(ws,row,1,ri,base,Font(bold=True,size=10,name='Calibri'),halign='center')
        _c(ws,row,2,v.get('original_text',''),base,Font(size=11,name='Calibri'))
        _c(ws,row,3,v.get('translated_text',''),base,Font(size=11,name='Calibri',color='1E40AF'))
        _c(ws,row,4,v.get('article_title',''),base,Font(size=9,name='Calibri',color='64748B'))
        _c(ws,row,5,'Mở bài ↗',base,link=v.get('article_url',''))
        ws.row_dimensions[row].height = 32
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()
